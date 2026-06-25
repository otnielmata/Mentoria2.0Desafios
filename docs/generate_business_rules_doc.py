from __future__ import annotations

import argparse
import json
import subprocess
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


TITLE = "Regras de Negocio do Sistema"
SUBTITLE = "Mentoria 2.0 Desafios"
REFERENCE_DATE = date(2026, 6, 25)

BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
TEXT = RGBColor(0x22, 0x22, 0x22)
MUTED = RGBColor(0x66, 0x66, 0x66)
LIGHT_FILL = "E8EEF5"
SOFT_FILL = "F6F8FB"
GRID = "D0D7E2"

STUDENT_MENU_ORDER = {
    "Inicio": 0,
    "Desafios": 1,
    "Check-list": 2,
    "Meu Perfil": 3,
    "Meus Grupos": 4,
    "Minha Pontuacao": 5,
    "Plano de Estudo": 6,
    "Ranking": 7,
}

ADMIN_MENU_ORDER = {
    "Dashboard": 0,
    "Alunos": 1,
    "Aprovacoes": 2,
    "Configuracoes": 3,
    "Cupons": 4,
    "Desafios": 5,
    "Pilares": 6,
    "Plano de Estudo": 7,
    "Ranking": 8,
    "Relatorios": 9,
    "Turmas": 10,
}


def normalize_label(value: str) -> str:
    return (
        value.replace("ç", "c")
        .replace("Ç", "C")
        .replace("ã", "a")
        .replace("Ã", "A")
        .replace("á", "a")
        .replace("Á", "A")
        .replace("é", "e")
        .replace("É", "E")
        .replace("í", "i")
        .replace("Í", "I")
        .replace("ó", "o")
        .replace("Ó", "O")
        .replace("ú", "u")
        .replace("Ú", "U")
        .replace("õ", "o")
        .replace("Õ", "O")
    )


def load_web_contract(project_root: Path, node_path: str) -> list[dict]:
    script = """
const contract = require('./web/src/contracts/api-endpoints.js');
process.stdout.write(JSON.stringify(contract.WEB_API_ENDPOINTS));
"""
    output = subprocess.check_output([node_path, "-e", script], cwd=project_root, text=True)
    return json.loads(output)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)

    for key, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_layout(table, widths_in: list[float], indent_dxa: int = 120) -> None:
    table.autofit = False
    table_pr = table._tbl.tblPr

    tbl_layout = table_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        table_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    tbl_w = table_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        table_pr.append(tbl_w)
    total_dxa = int(sum(widths_in) * 1440)
    tbl_w.set(qn("w:w"), str(total_dxa))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = table_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        table_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_grid = table._tbl.tblGrid
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths_in:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(int(width * 1440)))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths_in[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run_font(run, size: float, bold: bool = False, color: RGBColor = TEXT, italic: bool = False) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def style_paragraph(paragraph, after=6, before=0, line=1.25, alignment=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_after = Pt(after)
    fmt.space_before = Pt(before)
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = line
    paragraph.alignment = alignment


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = TEXT
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        style.paragraph_format.line_spacing = 1.25

    header = section.header.paragraphs[0]
    style_paragraph(header, after=0, line=1.0)
    header_run = header.add_run("Mentoria 2.0 Desafios | Guia consolidado de regras")
    set_run_font(header_run, 9.5, color=MUTED)

    footer = section.footer.paragraphs[0]
    style_paragraph(footer, after=0, line=1.0, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    footer_run = footer.add_run(f"Referencia do documento: {REFERENCE_DATE.isoformat()}")
    set_run_font(footer_run, 9.5, color=MUTED)


def add_cover(doc: Document) -> None:
    spacer = doc.add_paragraph()
    style_paragraph(spacer, after=0, line=1.0)
    spacer.paragraph_format.space_before = Pt(12)

    kicker = doc.add_paragraph()
    style_paragraph(kicker, after=8)
    kicker_run = kicker.add_run("DOCUMENTO DE REFERENCIA FUNCIONAL")
    set_run_font(kicker_run, 11, bold=True, color=MUTED)

    title = doc.add_paragraph()
    style_paragraph(title, after=4)
    title_run = title.add_run(TITLE)
    set_run_font(title_run, 23, bold=True, color=TEXT)

    subtitle = doc.add_paragraph()
    style_paragraph(subtitle, after=16)
    subtitle_run = subtitle.add_run(SUBTITLE)
    set_run_font(subtitle_run, 14, color=MUTED)

    metadata = doc.add_table(rows=5, cols=2)
    set_table_layout(metadata, [1.35, 5.15])
    rows = [
        ("Projeto", "Mentoria 2.0 Desafios"),
        ("Referencia", REFERENCE_DATE.strftime("%d/%m/%Y")),
        ("Perfis", "Publico, aluno, professor e admin"),
        ("Escopo", "Fluxos Web, contratos de API e regras de negocio implementadas no repositorio"),
        ("Fontes-base", "README, Swagger, contrato Web/API e servicos de dominio"),
    ]
    for row_index, (label, value) in enumerate(rows):
        row = metadata.rows[row_index]
        set_cell_shading(row.cells[0], SOFT_FILL)
        p_label = row.cells[0].paragraphs[0]
        style_paragraph(p_label, after=0)
        run_label = p_label.add_run(label)
        set_run_font(run_label, 10.5, bold=True, color=DARK_BLUE)

        p_value = row.cells[1].paragraphs[0]
        style_paragraph(p_value, after=0)
        run_value = p_value.add_run(value)
        set_run_font(run_value, 10.5)

    note = doc.add_paragraph()
    style_paragraph(note, before=16)
    note_run = note.add_run(
        "Este documento descreve o comportamento atual do sistema conforme a implementacao existente. "
        "Quando houver divergencia entre tela e API, a API continua sendo a fonte de verdade."
    )
    set_run_font(note_run, 10.5, color=TEXT)

    return None


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    run = paragraph.add_run(text)
    set_run_font(run, {1: 16, 2: 13, 3: 12}[level], bold=True, color={1: BLUE, 2: BLUE, 3: DARK_BLUE}[level])


def add_paragraph(doc_or_cell, text: str, size: float = 11, color: RGBColor = TEXT, bold: bool = False, after: int = 6) -> None:
    paragraph = doc_or_cell.add_paragraph() if hasattr(doc_or_cell, "add_paragraph") else doc_or_cell.paragraphs[0]
    style_paragraph(paragraph, after=after)
    run = paragraph.add_run(text)
    set_run_font(run, size, bold=bold, color=color)


def add_labeled_paragraph(doc_or_cell, label: str, text: str, after: int = 6) -> None:
    paragraph = doc_or_cell.add_paragraph() if hasattr(doc_or_cell, "add_paragraph") else doc_or_cell.paragraphs[0]
    style_paragraph(paragraph, after=after)
    label_run = paragraph.add_run(f"{label}: ")
    set_run_font(label_run, 10.8, bold=True, color=DARK_BLUE)
    text_run = paragraph.add_run(text)
    set_run_font(text_run, 10.8, color=TEXT)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_layout(table, widths)
    header_row = table.rows[0]
    for index, header in enumerate(headers):
        cell = header_row.cells[index]
        set_cell_shading(cell, LIGHT_FILL)
        paragraph = cell.paragraphs[0]
        style_paragraph(paragraph, after=0)
        run = paragraph.add_run(header)
        set_run_font(run, 10, bold=True, color=DARK_BLUE)

    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            paragraph = cells[index].paragraphs[0]
            style_paragraph(paragraph, after=0)
            run = paragraph.add_run(value)
            set_run_font(run, 10.2, color=TEXT)

    for row in table.rows[1:]:
        for cell in row.cells:
            set_cell_shading(cell, "FFFFFF")


def sort_endpoints(endpoints: list[dict], role: str) -> list[dict]:
    if role == "aluno":
        order = STUDENT_MENU_ORDER
    else:
        order = ADMIN_MENU_ORDER

    def key(item: dict):
        menu_order = order.get(normalize_label(item.get("menu", "")), 99)
        return (menu_order, item.get("screen", ""), item.get("menu", ""), item.get("method", ""), item.get("path", ""))

    return sorted(endpoints, key=key)


def endpoint_rows(endpoints: list[dict]) -> list[list[str]]:
    return [
        [
            f"{endpoint.get('menu', '-') } / {endpoint.get('screen', '-')}",
            endpoint.get("method", "-"),
            f"/api{endpoint.get('path', '')}",
            endpoint.get("feature", "-"),
        ]
        for endpoint in endpoints
    ]


def active_only(endpoints: list[dict]) -> list[dict]:
    return [endpoint for endpoint in endpoints if not endpoint.get("future")]


def filter_endpoints(endpoints: list[dict], profile: str) -> list[dict]:
    result = []
    for endpoint in endpoints:
        roles = endpoint.get("roles") or []
        if endpoint.get("public"):
            continue
        if profile == "aluno" and "aluno" in roles:
            result.append(endpoint)
        elif profile == "professor_admin" and ("professor" in roles or "admin" in roles) and roles != ["admin"]:
            result.append(endpoint)
        elif profile == "admin" and roles == ["admin"]:
            result.append(endpoint)
    return result


def add_public_section(doc: Document, endpoints: list[dict]) -> None:
    add_heading(doc, "5. Perfil Publico", 1)
    add_paragraph(
        doc,
        "O acesso publico fica restrito aos fluxos de entrada no sistema. Nao existe selecao manual de perfil na inscricao publica e a API cria sempre um usuario com papel de aluno e status ativo.",
    )

    add_heading(doc, "5.1 Web publica", 2)
    add_table(
        doc,
        ["Tela", "Finalidade", "Regra central"],
        [
            ["Login", "Autenticar usuario existente", "Sessao autenticada depende de credenciais validas e usuario ativo."],
            ["Registro", "Criar novo acesso publico", "O fluxo nao cria professor nem admin; o cadastro publico sempre gera aluno."],
        ],
        [1.1, 1.65, 3.75],
    )

    add_heading(doc, "5.2 API publica", 2)
    public_rows = endpoint_rows([endpoint for endpoint in active_only(endpoints) if endpoint.get("public")])
    add_table(doc, ["Tela / Menu", "Metodo", "Endpoint", "Funcionalidade"], public_rows, [1.55, 0.75, 2.25, 1.95])

    add_heading(doc, "5.3 Regras de negocio do perfil publico", 2)
    add_labeled_paragraph(doc, "Registro", "A rota publica de cadastro nao aceita governanca de perfil pelo usuario final. O objetivo e apenas criar alunos.")
    add_labeled_paragraph(doc, "Login", "Usuarios inativos nao podem acessar o sistema. O tratamento de autenticacao continua centralizado na API.")
    add_labeled_paragraph(doc, "Aliases", "A API ainda mantem aliases legados em /api/usuarios/registro e /api/usuarios/login, mas o contrato principal da Web usa /api/auth/register e /api/auth/login.")


def add_common_sections(doc: Document, future_endpoints: list[dict]) -> None:
    add_heading(doc, "1. Objetivo e escopo", 1)
    add_paragraph(
        doc,
        "Este documento consolida as regras de negocio atualmente implementadas no projeto Mentoria 2.0 Desafios, com separacao entre experiencia Web e contratos de API por perfil de usuario.",
    )
    add_labeled_paragraph(doc, "Fonte de verdade", "A API define autorizacao, pontuacao, ranking, validacao de cupons, auditoria e consolidacao de relatorios. A Web apenas consome e apresenta esses dados.")
    add_labeled_paragraph(doc, "Data de referencia", f"Todas as regras aqui descritas refletem o repositorio na referencia de {REFERENCE_DATE.strftime('%d/%m/%Y')}.")

    add_heading(doc, "2. Perfis e governanca de acesso", 1)
    add_table(
        doc,
        ["Perfil", "Web", "API", "Observacoes"],
        [
            ["Publico", "Login e registro", "Rotas publicas sem JWT", "Registro publico gera sempre aluno ativo."],
            ["Aluno", "Painel, desafios, checklist, grupos, pontuacao, perfil e plano de estudo", "Rotas com JWT e role aluno", "Pode atuar somente sobre os proprios dados e envios."],
            ["Professor", "Usa a mesma casca administrativa do admin, sem Configuracoes", "Rotas com JWT e role professor ou admin", "Atua em cadastros, aprovacoes, dashboards e relatorios."],
            ["Admin", "Mesmo menu administrativo do professor, com Configuracoes", "Todas as rotas de professor e as rotas exclusivas de usuarios", "Faz a gestao de usuarios e perfis do sistema."],
        ],
        [0.95, 2.1, 1.55, 1.9],
    )
    add_labeled_paragraph(doc, "Observacao de interface", "Na Web, o papel professor e tratado com a mesma shell visual do admin. A diferenca pratica e que somente o admin enxerga Configuracoes.")

    add_heading(doc, "3. Regras transversais do sistema", 1)
    add_labeled_paragraph(doc, "Contrato Web/API", "Toda tela da Web precisa consumir endpoint declarado no contrato central web/src/contracts/api-endpoints.js. Endpoints futuros devem ficar marcados com future e futureReason.")
    add_labeled_paragraph(doc, "Tratamento de erro", "Resposta 404 representa funcionalidade indisponivel ou divergencia de rota e nao deve derrubar a sessao. Apenas 401 deve disparar fluxo de autenticacao.")
    add_labeled_paragraph(doc, "Soft delete", "Alunos, usuarios, turmas e itens semelhantes sao inativados logicamente para preservar historico de envios, pontuacoes, auditoria e vinculos.")
    add_labeled_paragraph(doc, "Pilares fixos", "A API faz seed idempotente dos 7 pilares padrao do Metodo do Alavanque e usa esses pilares como base do cadastro de desafios e da distribuicao de pontos.")
    add_labeled_paragraph(doc, "Dashboards e relatorios", "A consolidacao acontece na API. A Web nao recalcula ranking, engajamento, baixa participacao, cupons nem pontuacao para compensar ausencia de endpoint.")
    add_labeled_paragraph(doc, "Pontuacao total do ecossistema", "O sistema considera pontuacoes geradas por desafios aprovados, pontos extras administrativos e pontos do checklist de planejamento para compor ranking, dashboard e cupons.")

    add_heading(doc, "3.1 Itens futuros mapeados no contrato", 2)
    future_rows = [
        [
            ", ".join(endpoint.get("roles", [])) or "n/d",
            endpoint.get("method", "-"),
            f"/api{endpoint.get('path', '')}",
            endpoint.get("futureReason", "Sem justificativa informada."),
        ]
        for endpoint in future_endpoints
    ]
    add_table(doc, ["Perfil", "Metodo", "Endpoint", "Motivo do futuro"], future_rows, [1.15, 0.75, 1.95, 2.65])


def add_study_plan_module(doc: Document, endpoints: list[dict]) -> None:
    add_heading(doc, "4. Modulo Plano de Estudo e Check-list", 1)
    add_paragraph(
        doc,
        "Este capitulo centraliza as regras do modulo Plano de Estudo para deixar o tema facil de localizar. Aqui entram agenda oficial da mentoria, planejamento pessoal, planejamento semanal replicado e Check-list.",
    )

    student_plan_endpoints = [
        endpoint
        for endpoint in active_only(filter_endpoints(endpoints, "aluno"))
        if endpoint.get("feature") == "Plano de Estudo"
    ]
    admin_plan_endpoints = [
        endpoint
        for endpoint in active_only(filter_endpoints(endpoints, "professor_admin"))
        if endpoint.get("feature") == "Plano de Estudo"
    ]

    add_heading(doc, "4.1 Visao geral do modulo", 2)
    add_table(
        doc,
        ["Perspectiva", "Objetivo", "Leitura principal"],
        [
            ["Aluno", "Organizar a propria rotina de estudo", "Calendario mensal, planejamento pessoal, replicacao semanal e Check-list de execucao."],
            ["Professor/Admin", "Organizar a agenda oficial da mentoria", "Cadastro de eventos por turma para orientar e sincronizar a jornada do aluno."],
            ["API", "Centralizar a regra operacional", "Validacao de datas, permissao, consolidacao da agenda e reflexo em pontuacao e cupons."],
        ],
        [1.2, 1.9, 3.4],
    )

    add_heading(doc, "4.2 Web do aluno no Plano de Estudo", 2)
    add_labeled_paragraph(doc, "Calendario", "O calendario ocupa a area principal da tela do aluno e combina eventos oficiais da mentoria com o planejamento pessoal.")
    add_labeled_paragraph(doc, "Novo planejamento", "O formulario de criacao fica na parte inferior para preservar a leitura do calendario.")
    add_labeled_paragraph(doc, "Eventos oficiais", "Eventos cadastrados pela mentoria aparecem como somente leitura. O aluno pode consultar os detalhes, mas nao pode editar nem excluir esses itens.")
    add_labeled_paragraph(doc, "Itens pessoais", "O aluno cria, edita e exclui apenas os proprios planejamentos.")
    add_labeled_paragraph(doc, "Check-list", "As tarefas planejadas alimentam diretamente o menu Check-list. Nao existe cadastro de tarefa separado do planejamento.")

    add_heading(doc, "4.3 API do aluno no Plano de Estudo", 2)
    add_table(doc, ["Tela / Menu", "Metodo", "Endpoint", "Funcionalidade"], endpoint_rows(sort_endpoints(student_plan_endpoints, "aluno")), [1.55, 0.75, 2.2, 1.95])

    add_heading(doc, "4.4 Regras de planejamento pessoal do aluno", 2)
    add_labeled_paragraph(doc, "Criacao", "Data de inicio e obrigatoria. Data final e opcional, mas quando informada nao pode ser anterior ao inicio.")
    add_labeled_paragraph(doc, "Persistencia", "Cada item guarda inicio, fim, data planejada e janela semanal de pontuacao, pois esses campos sao usados no Check-list e na regra de pontos.")
    add_labeled_paragraph(doc, "Atualizacao", "Ao alterar uma tarefa, a API reavalia a janela de pontuacao e resincroniza cupons quando a mudanca interfere na pontuacao do aluno.")
    add_labeled_paragraph(doc, "Exclusao", "A exclusao e logica. O item sai do uso ativo e a sincronizacao de cupons e refeita.")

    add_heading(doc, "4.5 Regras do planejamento semanal replicado", 2)
    add_labeled_paragraph(doc, "Acionamento", "Na Web, o aluno pode marcar a opcao de planejar a semana inteira.")
    add_labeled_paragraph(doc, "Duracao minima", "A replicacao garante pelo menos 90 minutos por dia. Se o intervalo informado for menor ou invalido, o sistema usa 90 minutos como base.")
    add_labeled_paragraph(doc, "Janela temporal", "A replicacao considera os proximos 7 dias a partir da data inicial escolhida.")
    add_labeled_paragraph(doc, "Bloqueio por evento ao vivo", "Dias com evento oficial do tipo ao vivo nao recebem copia automatica do estudo.")
    add_labeled_paragraph(doc, "Horario", "O horario do primeiro estudo e reaproveitado nos demais dias que forem liberados para criacao.")

    add_heading(doc, "4.6 Agenda oficial da mentoria", 2)
    add_labeled_paragraph(doc, "Tipos de evento", "Professor e admin trabalham com tres tipos oficiais: ao vivo, modulo gravado e conteudo especial.")
    add_labeled_paragraph(doc, "Papel no calendario do aluno", "Esses eventos entram na agenda consolidada do aluno e orientam a leitura da trilha.")
    add_labeled_paragraph(doc, "Influencia na replicacao", "Somente eventos ao vivo bloqueiam a criacao automatica do estudo semanal replicado.")

    add_heading(doc, "4.7 Check-list e pontuacao do planejamento", 2)
    add_labeled_paragraph(doc, "Liberacao do check", "O aluno so pode marcar o check em tarefas do dia atual ou de datas anteriores. Tarefas futuras ficam bloqueadas.")
    add_labeled_paragraph(doc, "Ocultacao de concluidas", "Quando uma tarefa e marcada, ela sai da lista de pendentes e vai para o filtro de concluidas. O aluno pode desfazer o check para devolver a tarefa aos pendentes.")
    add_labeled_paragraph(doc, "Pontuacao semanal", "A API consolida a pontuacao em janelas de 7 dias: 1 a 3 dias com check valem 1 ponto, 4 a 6 dias valem 2 pontos e 7 dias valem 3 pontos.")
    add_labeled_paragraph(doc, "Contagem por dia", "Mesmo com varias tarefas no mesmo dia, apenas um dia com check conta para a pontuacao daquela janela.")
    add_labeled_paragraph(doc, "Teto por janela", "O maximo conquistado em uma janela semanal e 3 pontos.")

    add_heading(doc, "4.8 Impacto do modulo em outros componentes", 2)
    add_labeled_paragraph(doc, "Minha Pontuacao", "Os pontos do Check-list entram no total de pontos do aluno e aparecem em separado no resumo da tela.")
    add_labeled_paragraph(doc, "Dashboard do aluno", "O painel do aluno incorpora o resumo do checklist, a pontuacao total e os reflexos disso em ranking e cupons.")
    add_labeled_paragraph(doc, "Dashboard administrativo", "O ranking administrativo tambem recebe os pontos do checklist quando a API consolida os alunos.")
    add_labeled_paragraph(doc, "Ranking e relatorios", "Quando o recorte nao filtra por pilar ou tipo, os pontos do checklist entram na classificacao e nos relatorios por aluno.")
    add_labeled_paragraph(doc, "Cupons", "A sincronizacao de cupons usa a soma da pontuacao tradicional com os pontos do checklist para definir quantos cupons o aluno deve possuir.")

    add_heading(doc, "4.9 API administrativa do Plano de Estudo", 2)
    add_table(doc, ["Tela / Menu", "Metodo", "Endpoint", "Funcionalidade"], endpoint_rows(sort_endpoints(admin_plan_endpoints, "admin")), [1.55, 0.75, 2.2, 1.95])
    add_labeled_paragraph(doc, "Governanca administrativa", "Professor e admin mantem a agenda oficial da mentoria; o aluno apenas consome esses eventos em modo leitura.")


def add_student_section(doc: Document, endpoints: list[dict]) -> None:
    add_heading(doc, "6. Perfil Aluno", 1)
    add_paragraph(
        doc,
        "O aluno concentra a execucao do metodo na propria trilha. A interface atual do aluno trabalha com o menu Início, Desafios, Check-list, Meu Perfil, Meus Grupos, Minha Pontuacao e Plano de Estudo.",
    )

    add_heading(doc, "6.1 Web do aluno", 2)
    add_table(
        doc,
        ["Menu", "Papel na experiencia", "Regras principais"],
        [
            ["Inicio", "Painel consolidado do aluno", "Exibe pontuacao total, posicao no ranking, desafios, cupons e numeros da sorte ja distribuidos."],
            ["Desafios", "Inscricao e envio", "O aluno se inscreve primeiro e so depois envia a entrega do desafio vinculado ao grupo."],
            ["Check-list", "Execucao do planejamento", "Mostra tarefas planejadas do aluno e controla o check apenas em tarefas de hoje ou datas passadas."],
            ["Meu Perfil", "Autogestao de dados basicos", "O aluno pode alterar nome e senha, mas nao altera perfil, status nem turma."],
            ["Meus Grupos", "Acompanhamento do grupo automatico", "Permite ver participantes e atualizar o link de contato do grupo."],
            ["Minha Pontuacao", "Historico e consolidado de pontos", "O total inclui desafios aprovados e pontos do checklist."],
            ["Plano de Estudo", "Agenda mensal e planejamento pessoal", "Mistura eventos oficiais da mentoria com o planejamento pessoal do aluno."],
        ],
        [1.1, 1.85, 3.55],
    )

    add_heading(doc, "6.2 API principal do aluno", 2)
    aluno_rows = endpoint_rows(sort_endpoints(active_only(filter_endpoints(endpoints, "aluno")), "aluno"))
    add_table(doc, ["Tela / Menu", "Metodo", "Endpoint", "Funcionalidade"], aluno_rows, [1.55, 0.75, 2.2, 1.95])
    add_labeled_paragraph(doc, "Observacao", "A API tambem expoe GET /api/ranking para aluno autenticado, embora a interface atual do aluno nao tenha um item de menu dedicado para ranking.")

    add_heading(doc, "6.3 Regras de negocio por modulo do aluno", 2)

    add_heading(doc, "6.3.1 Inicio, ranking e cupons visiveis ao aluno", 3)
    add_labeled_paragraph(doc, "Web", "O Inicio concentra os principais indicadores do aluno: pontuacao total, posicao, quantidade de desafios, distribuicao por pilar, resumo do checklist e bloco de cupons.")
    add_labeled_paragraph(doc, "API", "GET /api/dashboard/aluno consolida o painel e devolve os dados de ranking, pontuacao, checklist e cupons.")
    add_labeled_paragraph(doc, "Regra de pontuacao", "A pontuacao total do painel soma pontos de desafios aprovados e pontos obtidos no checklist de planejamento.")
    add_labeled_paragraph(doc, "Regra de cupons", "O painel mostra total de cupons, cupons pendentes, cupons validados, cupons aguardando numero da sorte e a lista de numeros da sorte ja distribuidos para o aluno.")

    add_heading(doc, "6.3.2 Desafios, inscricoes e envios", 3)
    add_labeled_paragraph(doc, "Web", "O modulo Desafios cobre inscricao em desafio, historico de envios do grupo do aluno e atualizacao do proprio envio enquanto ele ainda permite edicao.")
    add_labeled_paragraph(doc, "API", "GET /api/desafios, GET /api/desafios/inscricoes/minhas, POST /api/desafios/:id/inscricoes, GET /api/envios-desafios/meus, POST /api/envios-desafios e PATCH /api/envios-desafios/:id.")
    add_labeled_paragraph(doc, "Inscricao obrigatoria", "O aluno precisa estar inscrito no desafio antes de registrar uma entrega. A inscricao cria ou completa automaticamente um grupo compativel.")
    add_labeled_paragraph(doc, "Agrupamento", "Os grupos respeitam desafio, turma e modalidade. Modalidades normal e ingles nao se misturam.")
    add_labeled_paragraph(doc, "Prazo", "A data limite do desafio e inclusiva. Apos o fim do dia cadastrado, o desafio deixa de aceitar novas inscricoes.")
    add_labeled_paragraph(doc, "Envio", "Descricao e obrigatoria. Evidencia ou link e anexo sao opcionais. Ao enviar, o grupo inteiro passa a enxergar aquele envio.")
    add_labeled_paragraph(doc, "Edicao", "Envios podem ser alterados apenas enquanto estiverem pendentes ou em ajuste. Envio aprovado fica bloqueado para edicao.")

    add_heading(doc, "6.3.3 Integracao do aluno com o modulo Plano de Estudo", 3)
    add_labeled_paragraph(doc, "Referencia principal", "O detalhamento completo de Plano de Estudo e Check-list foi centralizado no capitulo 4 deste documento.")
    add_labeled_paragraph(doc, "Papel no ecossistema do aluno", "Esse modulo e o ponto de origem do Check-list, dos pontos de planejamento e da parcela de pontuacao que influencia ranking, dashboard e cupons.")

    add_heading(doc, "6.3.4 Meus Grupos, Minha Pontuacao e Meu Perfil", 3)
    add_labeled_paragraph(doc, "Meus Grupos", "O aluno enxerga os participantes do grupo automatico e qualquer integrante pode atualizar o canal de contato do grupo, como WhatsApp, Telegram ou Discord.")
    add_labeled_paragraph(doc, "Minha Pontuacao", "A tela lista consolidado de pontos, pontos por pilar, historico de lancamentos aprovados e resumo do checklist. Filtros por periodo continuam na API.")
    add_labeled_paragraph(doc, "Meu Perfil", "O aluno pode alterar nome e senha. Campos de perfil, status e turma ficam apenas para leitura no front e devem ser administrados fora deste modulo.")


def add_admin_section(doc: Document, endpoints: list[dict]) -> None:
    add_heading(doc, "7. Perfil Professor e Admin", 1)
    add_paragraph(
        doc,
        "Professor e admin compartilham o mesmo ambiente operacional da mentoria. O admin acumula tudo o que o professor faz e ainda assume a gestao global de usuarios em Configuracoes.",
    )

    add_heading(doc, "7.1 Web administrativa", 2)
    add_table(
        doc,
        ["Menu", "Papel operacional", "Regras principais"],
        [
            ["Dashboard", "Painel consolidado da mentoria", "Mostra alunos ativos, envios, aprovacoes pendentes, ranking, engajamento e cupons gerados."],
            ["Alunos", "Cadastro e manutencao de alunos", "Aceita cadastro manual, importacao CSV, edicao e inativacao logica."],
            ["Aprovacoes", "Avaliacao de envios", "Permite aprovar, reprovar, pedir ajuste e lancar pontos extras."],
            ["Cupons", "Governanca de cupons e numeros da sorte", "Lista cupons por aluno, lista global validada e aciona a distribuicao dos numeros."],
            ["Desafios", "Cadastro da trilha executavel", "Controla pilares, pontos fixos, recorrencia, certificados e tamanho de grupo."],
            ["Pilares", "Manutencao dos pilares do metodo", "Aceita cadastro, edicao e inativacao de pilares."],
            ["Plano de Estudo", "Agenda oficial da mentoria", "Cadastra eventos por turma com tipos ao vivo, modulo gravado e conteudo especial."],
            ["Ranking", "Consulta ranqueada da base", "Lista alunos por pontuacao total e aceita filtros administrativos."],
            ["Relatorios", "Leitura analitica e gerencial", "Reune participacao, pilares, grupos, numeros da sorte e baixa participacao."],
            ["Turmas", "Organizacao das turmas", "Mantem turmas, historico de alunos e estado da turma."],
        ],
        [1.1, 1.8, 3.6],
    )
    add_labeled_paragraph(doc, "Menu exclusivo do admin", "Configuracoes aparece apenas para admin e concentra a gestao de usuarios e perfis.")

    add_heading(doc, "7.2 API principal do professor/admin", 2)
    admin_rows = endpoint_rows(sort_endpoints(active_only(filter_endpoints(endpoints, "professor_admin")), "admin"))
    add_table(doc, ["Tela / Menu", "Metodo", "Endpoint", "Funcionalidade"], admin_rows, [1.55, 0.75, 2.2, 1.95])
    add_labeled_paragraph(doc, "Observacao", "A API tambem mantem aliases administrativos em /api/admin/* e algumas rotas complementares de detalhe, mas a tabela acima prioriza o contrato principal consumido pela Web.")

    add_heading(doc, "7.3 Regras de negocio por modulo administrativo", 2)

    add_heading(doc, "7.3.1 Dashboard, engajamento e ranking", 3)
    add_labeled_paragraph(doc, "Web", "O Dashboard administrativo mostra indicadores gerais, top ranking, ranking por turma e por pilar, desafios por pilar, cupons gerados e metricas de engajamento.")
    add_labeled_paragraph(doc, "API", "GET /api/dashboard/admin consolida essas leituras.")
    add_labeled_paragraph(doc, "Regra de ranking", "O ranking administrativo soma pontos aprovados e pontos de checklist. Empates usam nome e id como criterio de ordenacao compartilhada.")
    add_labeled_paragraph(doc, "Regra de escopo", "A API pode ocultar alunos inativos do ranking e dos paines, preservando o historico mas retirando quem nao deve mais competir.")
    add_labeled_paragraph(doc, "Indicador de cupons", "O dashboard devolve cuponsGerados como total de cupons ativos gerados na base, independentemente de validacao ou numero da sorte.")

    add_heading(doc, "7.3.2 Cupons, validacao e numeros da sorte", 3)
    add_labeled_paragraph(doc, "Web", "O menu Cupons lista cada aluno com total de cupons, pendencias, validacoes e numeros da sorte; tambem exibe uma listagem global cupom por cupom.")
    add_labeled_paragraph(doc, "API", "GET /api/cupons/alunos, GET /api/cupons/numeros-sorte, POST /api/cupons/distribuicao/numeros-sorte e GET /api/relatorios/cupons-sorte.")
    add_labeled_paragraph(doc, "Geracao", "O sistema gera 1 cupom a cada 10 pontos acumulados pelo aluno. Para esse calculo entram pontuacoes aprovadas, pontos extras e pontos do checklist.")
    add_labeled_paragraph(doc, "Sincronizacao", "Os cupons acompanham a pontuacao vigente. Se a pontuacao cair por reabertura de checklist ou cancelamento de itens, cupons excedentes podem ser cancelados.")
    add_labeled_paragraph(doc, "Validacao", "Cupons nascem pendentes. Eles so viram validados quando um envio aprovado pertence a um desafio marcado com certificado postado.")
    add_labeled_paragraph(doc, "Distribuicao de numeros", "Somente cupons validados participam da distribuicao. O botao administrativo preserva numeros ja distribuidos e atribui novos numeros apenas aos cupons validados que ainda nao tinham numeracao.")
    add_labeled_paragraph(doc, "Ordenacao da distribuicao", "A numeracao segue a ordem cronologica de validacao e conquista do cupom, mantendo consistencia temporal para a listagem de numeros da sorte.")
    add_labeled_paragraph(doc, "Visibilidade", "O aluno visualiza seus numeros da sorte no Inicio, enquanto o administrador visualiza tanto o resumo por aluno quanto a listagem global dos cupons validados.")

    add_heading(doc, "7.3.3 Alunos e turmas", 3)
    add_labeled_paragraph(doc, "Web", "Alunos cobre cadastro manual, importacao por CSV, edicao e inativacao. Turmas cobre cadastro, manutencao e organizacao da base por turma.")
    add_labeled_paragraph(doc, "API", "GET/POST/PATCH/DELETE /api/alunos, POST /api/alunos/importar e GET/POST /api/turmas.")
    add_labeled_paragraph(doc, "CSV", "A importacao de alunos exige as colunas Nome, E-mail, Senha Inicial e Turma. A turma pode ser encontrada por nome, codigo ou id.")
    add_labeled_paragraph(doc, "Dados administrativos do aluno", "Professor e admin podem editar nome, email, senha, status, turma e o marcador administrativo discordJoined.")
    add_labeled_paragraph(doc, "Inativacao", "Exclusao de aluno e feita por soft delete. O historico operacional do aluno permanece preservado.")

    add_heading(doc, "7.3.4 Pilares e desafios", 3)
    add_labeled_paragraph(doc, "Web", "Pilares mantem a taxonomia do metodo. Desafios mantem a trilha executavel com pontuacao, prazo, recorrencia e requisitos de grupo.")
    add_labeled_paragraph(doc, "API", "GET/POST/PATCH/DELETE /api/pilares e GET/POST/PATCH/DELETE /api/desafios.")
    add_labeled_paragraph(doc, "Pilares", "Todo desafio precisa ter ao menos um pilar ativo com pontuacao positiva. O sistema nao aceita pilar repetido no mesmo desafio.")
    add_labeled_paragraph(doc, "Pontuacao fixa", "A soma dos pontos por pilar define a pontuacao base do desafio. O campo difficulty pode permanecer como metadado, mas nao substitui a pontuacao fixa.")
    add_labeled_paragraph(doc, "Tipos e grupos", "O desafio aceita tipo individual, grupo ou ambos. Para desafios com grupos, maxParticipantes deve ficar entre 1 e 5.")
    add_labeled_paragraph(doc, "Prazo", "Delivery date e inclusiva. Depois do fim do dia configurado, o desafio e inativado automaticamente e deixa de aceitar inscricoes.")
    add_labeled_paragraph(doc, "Pontos extras de apresentacao", "livePresentationPoints pode ser zero ou positivo e so entra no calculo quando a aprovacao marca apresentacao ao vivo.")
    add_labeled_paragraph(doc, "Certificado postado", "Quando o campo certificado postado fica ligado, a aprovacao desse desafio valida os cupons pendentes dos alunos pontuados naquele envio.")
    add_labeled_paragraph(doc, "Recorrencia", "Desafios recorrentes usam periodo diario, semanal ou mensal e bloqueiam aprovacao quando a soma do aluno no periodo excede o limite configurado.")

    add_heading(doc, "7.3.5 Integracao administrativa com o modulo Plano de Estudo", 3)
    add_labeled_paragraph(doc, "Referencia principal", "O detalhamento completo da agenda oficial da mentoria foi centralizado no capitulo 4 deste documento.")
    add_labeled_paragraph(doc, "Papel administrativo", "Professor e admin mantem os eventos que orientam o calendario do aluno e interferem na replicacao automatica de estudos.")

    add_heading(doc, "7.3.6 Aprovacoes, pontuacao e pontos extras", 3)
    add_labeled_paragraph(doc, "Web", "A fila de aprovacoes exibe envio, desafio, turma, participantes, evidencias, anexos e historico de status. O professor ou admin decide o resultado.")
    add_labeled_paragraph(doc, "API", "GET /api/envios-desafios/aprovacoes, PATCH /api/envios-desafios/aprovacoes e POST /api/pontuacoes/extras.")
    add_labeled_paragraph(doc, "Decisoes", "Aprovado, reprovado e ajuste sao as unicas decisoes validas. Feedback e obrigatorio em reprovacao e ajuste.")
    add_labeled_paragraph(doc, "Bloqueios", "Envio cancelado nao pode ser avaliado e envio aprovado nao pode receber nova decisao.")
    add_labeled_paragraph(doc, "Pontuacao automatica", "Ao aprovar, a API gera pontuacao para o aluno responsavel e para todos os participantes ativos do grupo.")
    add_labeled_paragraph(doc, "Protecao contra duplicidade", "A aprovacao verifica se a mesma evidencia ja gerou pontuacao para aquele desafio e impede pontuacao duplicada.")
    add_labeled_paragraph(doc, "Recorrencia", "Antes de gerar pontos, a aprovacao tambem verifica se o aluno excederia o limite de pontos do desafio recorrente no periodo.")
    add_labeled_paragraph(doc, "Pontos extras", "Pontos extras so podem ser concedidos para aluno ativo e pilar ativo. Depois do lancamento, a logica de cupons e resincronizada.")

    add_heading(doc, "7.3.7 Ranking, relatorios e auditoria", 3)
    add_labeled_paragraph(doc, "Web", "Ranking e Relatorios entregam leitura gerencial da operacao. Auditoria entra como leitura de rastreabilidade para perfis administrativos.")
    add_labeled_paragraph(doc, "API", "GET /api/ranking/admin, GET /api/relatorios/participacao, GET /api/relatorios/alunos/pilares, GET /api/relatorios/grupos-desafios, GET /api/relatorios/cupons-sorte, GET /api/admin/relatorios/baixa-participacao e GET /api/auditorias.")
    add_labeled_paragraph(doc, "Checklist nos relatorios", "Relatorios por aluno e ranking incorporam pontos do checklist quando o recorte nao esta filtrado por pilar ou tipo, mantendo coerencia com o painel e a pontuacao total.")
    add_labeled_paragraph(doc, "Participacao", "Os relatorios de participacao consolidam totais por status, pontos aprovados, leitura por aluno, leitura por turma e baixa participacao.")
    add_labeled_paragraph(doc, "Rastreabilidade", "Criacao de envio, avaliacao e geracao de pontuacao registram eventos de auditoria sem expor senha, token, hash ou segredos.")


def add_admin_only_section(doc: Document, endpoints: list[dict]) -> None:
    add_heading(doc, "8. Perfil Admin exclusivo", 1)
    add_paragraph(
        doc,
        "O admin possui tudo o que o professor faz e ainda controla a governanca de usuarios do sistema. Esse bloco representa a area que nao e compartilhada com o professor.",
    )

    add_heading(doc, "8.1 Web exclusiva do admin", 2)
    add_table(
        doc,
        ["Menu", "Objetivo", "Regra central"],
        [
            ["Configuracoes", "Gestao de usuarios e perfis", "Permite criar, listar, editar e inativar usuarios aluno, professor ou admin."],
        ],
        [1.25, 1.9, 3.35],
    )

    add_heading(doc, "8.2 API exclusiva do admin", 2)
    admin_only_rows = endpoint_rows(sort_endpoints(active_only(filter_endpoints(endpoints, "admin")), "admin"))
    add_table(doc, ["Tela / Menu", "Metodo", "Endpoint", "Funcionalidade"], admin_only_rows, [1.55, 0.75, 2.2, 1.95])

    add_heading(doc, "8.3 Regras de negocio do admin exclusivo", 2)
    add_labeled_paragraph(doc, "Perfis", "Somente o admin cria ou altera usuarios com papeis aluno, professor ou admin pelo modulo Configuracoes.")
    add_labeled_paragraph(doc, "Soft delete", "Exclusao de usuario e logica. O usuario deixa de operar, mas o historico de auditoria e de dominio continua preservado.")
    add_labeled_paragraph(doc, "Autoprotecao", "O administrador autenticado nao pode excluir a propria conta.")
    add_labeled_paragraph(doc, "Configuracoes funcionais", "A API ainda expoe GET /api/configuracoes como leitura tecnica para professor ou admin, mas esse contrato nao representa um modulo de edicao seguro na Web atual.")


def add_appendix(doc: Document) -> None:
    add_heading(doc, "9. Observacoes finais e fontes consideradas", 1)
    add_labeled_paragraph(doc, "Fontes estruturais", "README.md, docs/swagger.yaml, docs/contrato-web-api.md e web/src/contracts/api-endpoints.js.")
    add_labeled_paragraph(doc, "Fontes de regra atualizada", "Servicos de dominio como plano-estudo.service.js, cupom.service.js, pontuacao.service.js, desafio.service.js, ranking.service.js, me-dashboard.service.js e admin-dashboard.service.js.")
    add_labeled_paragraph(doc, "Limite da leitura", "O documento foi redigido para refletir comportamento implementado. Alias tecnicos e rotas auxiliares que nao mudam a regra principal foram consolidados em notas e nao em inventario exaustivo.")


def build_document(project_root: Path, output_path: Path, node_path: str) -> None:
    endpoints = load_web_contract(project_root, node_path)
    future_endpoints = [endpoint for endpoint in endpoints if endpoint.get("future")]

    doc = Document()
    configure_styles(doc)
    add_cover(doc)
    add_common_sections(doc, future_endpoints)
    add_study_plan_module(doc, endpoints)
    add_public_section(doc, endpoints)
    add_student_section(doc, endpoints)
    add_admin_section(doc, endpoints)
    add_admin_only_section(doc, endpoints)
    add_appendix(doc)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def main() -> None:
    parser = argparse.ArgumentParser(description="Gera o documento DOCX de regras de negocio do projeto.")
    parser.add_argument("--project-root", required=True)
    parser.add_argument("--output", required=True)
    parser.add_argument("--node", required=True)
    args = parser.parse_args()

    build_document(Path(args.project_root), Path(args.output), args.node)


if __name__ == "__main__":
    main()
